from docx import Document
from docx.shared import Inches


class EvidenceGenerator():
    def __init__(self, project, execTime, finalResult):
        self.document = Document()
        self.fillHeader(project, execTime, finalResult)
        
    def fillHeader(self, project, execTime, finalResult):
        self.document.add_heading('Evidence document', 0)
        records = (
        (project, execTime, finalResult),
    )   
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Project'
        hdr_cells[1].text = 'Execution time'
        hdr_cells[2].text = 'Final Result'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc
 
    def addEvidence(self, testName, evidenceName, picture):
        self.document.add_heading(testName, level=1)
        self.document.add_paragraph(evidenceName, style='Intense Quote')
        self.document.add_picture(picture, width=Inches(6.25))
        self.document.add_page_break()

    def createDocument(self, doc_test):
        self.document.save(doc_test)
